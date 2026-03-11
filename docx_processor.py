#!/usr/bin/env python3
"""
Модуль для обработки DOCX файлов
"""

import sys
import os
from typing import Dict, List, Optional

class DocxProcessor:
    """Обработчик DOCX файлов"""
    
    def __init__(self):
        self.docx_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Проверить доступность библиотек"""
        try:
            import docx
            self.docx_available = True
            self.docx_module = docx
        except ImportError:
            self.docx_available = False
    
    def extract_text(self, file_path: str) -> Dict:
        """
        Извлечь текст из DOCX файла
        
        Args:
            file_path: Путь к DOCX файлу
            
        Returns:
            Словарь с результатом
        """
        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": f"Файл не найден: {file_path}",
                "text": ""
            }
        
        if not self.docx_available:
            return {
                "success": False,
                "error": "Библиотека python-docx не установлена",
                "text": "",
                "install_command": "pip3 install python-docx"
            }
        
        try:
            from docx import Document
            
            # Открыть документ
            doc = Document(file_path)
            
            # Извлечь весь текст
            full_text = []
            
            # Параграфы
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Метаданные
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            return {
                "success": True,
                "text": "\n".join(full_text),
                "metadata": metadata,
                "file_path": file_path,
                "file_size": os.path.getsize(file_path)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Ошибка при обработке DOCX: {str(e)}",
                "text": ""
            }
    
    def get_info(self, file_path: str) -> Dict:
        """Получить информацию о DOCX файле"""
        if not os.path.exists(file_path):
            return {"error": "Файл не найден"}
        
        try:
            import zipfile
            from datetime import datetime
            
            info = {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "modified": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
                "is_docx": False
            }
            
            # Проверить, что это ZIP архив (DOCX основан на ZIP)
            if zipfile.is_zipfile(file_path):
                info["is_docx"] = True
                with zipfile.ZipFile(file_path, 'r') as zf:
                    info["files_in_archive"] = zf.namelist()
                    info["is_valid_docx"] = any(f.startswith('word/') for f in zf.namelist())
            
            return info
            
        except Exception as e:
            return {"error": str(e)}
    
    def convert_to_txt(self, file_path: str, output_path: Optional[str] = None) -> Dict:
        """Конвертировать DOCX в TXT"""
        result = self.extract_text(file_path)
        
        if not result["success"]:
            return result
        
        if output_path is None:
            output_path = file_path.replace('.docx', '.txt').replace('.doc', '.txt')
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result["text"])
            
            return {
                "success": True,
                "input": file_path,
                "output": output_path,
                "text_length": len(result["text"]),
                "output_size": os.path.getsize(output_path)
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Ошибка при сохранении TXT: {str(e)}"
            }


def main():
    """Тестовая функция"""
    processor = DocxProcessor()
    
    print("=== DOCX Processor ===")
    print(f"python-docx доступен: {processor.docx_available}")
    
    if not processor.docx_available:
        print("\nДля установки выполните:")
        print("  pip3 install python-docx")
        print("  или")
        print("  apt-get install python3-docx")
    
    # Пример использования
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        if os.path.exists(file_path):
            print(f"\nОбработка файла: {file_path}")
            info = processor.get_info(file_path)
            print(f"Информация: {info}")
            
            if info.get("is_docx", False):
                result = processor.extract_text(file_path)
                if result["success"]:
                    print(f"\nТекст ({len(result['text'])} символов):")
                    print("-" * 50)
                    print(result["text"][:500] + "..." if len(result["text"]) > 500 else result["text"])
                    print("-" * 50)
                else:
                    print(f"Ошибка: {result['error']}")
        else:
            print(f"Файл не найден: {file_path}")


if __name__ == "__main__":
    main()